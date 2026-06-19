"""Backend tests for P1 — per-file download access control + in-browser preview.

Covers the API surface that powers the "View-only by default + Allow download
per-file" institutional VDR access pattern (iter-22).
"""

import io
import os
import requests
import pytest


def _load_backend_url():
    v = os.environ.get("REACT_APP_BACKEND_URL")
    if v:
        return v.rstrip("/")
    with open("/app/frontend/.env") as fh:
        for line in fh:
            if line.startswith("REACT_APP_BACKEND_URL="):
                return line.split("=", 1)[1].strip().rstrip("/")
    raise RuntimeError("REACT_APP_BACKEND_URL not configured")


BASE_URL = _load_backend_url()
API = f"{BASE_URL}/api"

SELLER = ("mira@workz.example.com", "WorkzPass123!")
BUYER = ("alex@workz.example.com", "WorkzPass123!")


def _login(email, password):
    r = requests.post(f"{API}/auth/login", json={"email": email, "password": password}, timeout=30)
    assert r.status_code == 200
    return r.json()["token"]


def _hdr(tok):
    return {"Authorization": f"Bearer {tok}"}


@pytest.fixture(scope="module")
def seller_tok():
    return _login(*SELLER)


@pytest.fixture(scope="module")
def buyer_tok():
    return _login(*BUYER)


@pytest.fixture(scope="module")
def active_room(seller_tok):
    """Pick any active vault where Mira is the literal seller."""
    me = requests.get(f"{API}/auth/me", headers=_hdr(seller_tok), timeout=30).json()
    rooms = requests.get(f"{API}/deal-rooms", headers=_hdr(seller_tok), timeout=30).json()
    own = [r for r in rooms if r.get("seller_id") == me["id"] and r.get("status") == "active"]
    assert own, "No active vault for seller"
    return own[0]["id"]


@pytest.fixture()
def uploaded_file(seller_tok, active_room):
    """Upload a tiny text file to the room, return its id; clean up via room view."""
    body = b"TEST iter22 access-control payload\n" * 20
    r = requests.post(
        f"{API}/deal-rooms/{active_room}/files/binary",
        files={"file": ("TEST_iter22_access.txt", io.BytesIO(body), "text/plain")},
        data={"folder": "other"},
        headers=_hdr(seller_tok), timeout=60,
    )
    assert r.status_code == 200, r.text
    return r.json()["id"], body


class TestDownloadAccessControl:
    def test_default_is_view_only_for_buyer(self, buyer_tok, active_room, uploaded_file):
        fid, _body = uploaded_file
        r = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/download",
                         headers=_hdr(buyer_tok), timeout=30)
        assert r.status_code == 403, "buyer must be blocked by default"
        assert "view-only" in r.text.lower() or "view only" in r.text.lower()

    def test_seller_always_downloads(self, seller_tok, active_room, uploaded_file):
        fid, body = uploaded_file
        r = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/download",
                         headers=_hdr(seller_tok), timeout=30)
        assert r.status_code == 200, r.text
        assert r.content == body, "seller download must return verbatim plaintext"

    def test_seller_toggles_and_buyer_downloads(self, seller_tok, buyer_tok, active_room, uploaded_file):
        fid, body = uploaded_file
        # Flip download_allowed=True
        r = requests.patch(
            f"{API}/deal-rooms/{active_room}/files/{fid}/access",
            json={"download_allowed": True},
            headers=_hdr(seller_tok), timeout=30,
        )
        assert r.status_code == 200, r.text
        assert r.json()["download_allowed"] is True
        # Buyer can now download
        rdl = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/download",
                           headers=_hdr(buyer_tok), timeout=30)
        assert rdl.status_code == 200, rdl.text
        assert rdl.content == body
        # Flip it back
        r2 = requests.patch(
            f"{API}/deal-rooms/{active_room}/files/{fid}/access",
            json={"download_allowed": False},
            headers=_hdr(seller_tok), timeout=30,
        )
        assert r2.status_code == 200
        rdl2 = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/download",
                            headers=_hdr(buyer_tok), timeout=30)
        assert rdl2.status_code == 403

    def test_buyer_cannot_set_access_policy(self, buyer_tok, active_room, uploaded_file):
        fid, _body = uploaded_file
        r = requests.patch(
            f"{API}/deal-rooms/{active_room}/files/{fid}/access",
            json={"download_allowed": True},
            headers=_hdr(buyer_tok), timeout=30,
        )
        assert r.status_code == 403, "buyer must not be able to flip access policy"


class TestPreviewEndpoint:
    def test_preview_returns_inline_for_buyer_regardless_of_download_flag(
        self, buyer_tok, active_room, uploaded_file,
    ):
        """Preview is always available to participants — that's the whole point
        of the view-only model. Buyer must be able to preview even when
        download_allowed=False."""
        fid, _body = uploaded_file
        r = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/preview",
                         headers=_hdr(buyer_tok), timeout=60)
        assert r.status_code == 200, r.text
        assert "inline" in (r.headers.get("Content-Disposition") or "").lower()

    def test_preview_logs_audit_event(self, buyer_tok, seller_tok, active_room, uploaded_file):
        fid, _body = uploaded_file
        # Hit preview to generate an event
        requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/preview",
                     headers=_hdr(buyer_tok), timeout=60)
        # Read the Activity feed and look for dealroom.file.preview entries
        act = requests.get(f"{API}/deal-rooms/{active_room}/activity",
                           headers=_hdr(seller_tok), timeout=30).json()
        preview_events = [
            e for e in act["events"]
            if e["action"] == "dealroom.file.preview" and (e.get("meta") or {}).get("file_id") == fid
        ]
        assert preview_events, "preview should log a dealroom.file.preview audit event"

    def test_preview_office_conversion_when_libreoffice_present(
        self, seller_tok, active_room,
    ):
        """When LibreOffice is installed, DOCX previews come back as application/pdf."""
        import shutil as _sh
        if not (_sh.which("soffice") or _sh.which("libreoffice")):
            pytest.skip("LibreOffice not installed in this environment")
        try:
            from docx import Document
        except ImportError:
            pytest.skip("python-docx not available")
        buf = io.BytesIO()
        doc = Document()
        doc.add_heading("iter22 office preview test", 0)
        doc.add_paragraph("Should round-trip through LibreOffice to PDF.")
        doc.save(buf)
        buf.seek(0)
        up = requests.post(
            f"{API}/deal-rooms/{active_room}/files/binary",
            files={"file": ("iter22_test.docx", buf,
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            data={"folder": "other"},
            headers=_hdr(seller_tok), timeout=60,
        )
        assert up.status_code == 200, up.text
        fid = up.json()["id"]
        r = requests.get(f"{API}/deal-rooms/{active_room}/files/{fid}/preview",
                         headers=_hdr(seller_tok), timeout=120)
        assert r.status_code == 200, r.text
        assert r.headers.get("content-type", "").startswith("application/pdf"), \
            f"DOCX preview must be PDF, got {r.headers.get('content-type')}"
        assert r.content.startswith(b"%PDF-"), "PDF magic bytes missing"
